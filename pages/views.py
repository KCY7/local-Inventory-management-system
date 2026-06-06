import subprocess
import os
import io
import uuid
import re 
from pathlib import Path
from docx.shared import Pt

from django.core.paginator import Paginator
from docxtpl import DocxTemplate
from django.db.models import Q
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout
from django.utils import timezone
from django.http import HttpResponse, JsonResponse 
from django.conf import settings
from django.contrib import messages

# --- IMPORTS FOR LOGIN ---
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import login as auth_login

# --- MODELS & FORMS ---
from .models import Room, PersonalComputer, Laptop, Tool, BorrowedItem, TransactionHistory
from .forms import PCForm, ToolForm, LaptopForm

# =========================================================
#  GLOBAL HELPERS
# =========================================================

def filter_blank_slots(item_list):
    """
    Blank slot filter: Removes None, empty strings, and whitespace-only items from a list.
    Ensures empty form inputs don't trigger logic loops.
    """
    return [item for item in item_list if item and str(item).strip()]

def is_duplicate_sn(sn, exclude_tool_id=None, exclude_pc_id=None, allow_loose=False):
    if not sn or not str(sn).strip():
        return False
    
    clean_sn = str(sn).strip()
    
    # 1. Check standalone Tools
    tool_qs = Tool.objects.filter(serial_number__iexact=clean_sn)
    if exclude_tool_id:
        tool_qs = tool_qs.exclude(id=exclude_tool_id)
    if exclude_pc_id:
        tool_qs = tool_qs.exclude(linked_pc_id=exclude_pc_id)
        
    # THE FIX: If we are adding/updating a PC, ignore loose tools in storage so we can fetch them!
    if allow_loose:
        tool_qs = tool_qs.exclude(linked_pc__isnull=True)
        
    if tool_qs.exists():
        return True
        
    # 2. Check PC Components
    pc_qs = PersonalComputer.objects.filter(
        Q(processor_sn__iexact=clean_sn) |
        Q(ram_sn__iexact=clean_sn) | Q(ram_2_sn__iexact=clean_sn) | Q(ram_3_sn__iexact=clean_sn) | Q(ram_4_sn__iexact=clean_sn) |
        Q(storage_sn__iexact=clean_sn) | Q(storage_2_sn__iexact=clean_sn) | Q(storage_3_sn__iexact=clean_sn) | Q(storage_4_sn__iexact=clean_sn) |
        Q(graphics_card_sn__iexact=clean_sn) |
        Q(motherboard_sn__iexact=clean_sn) |
        Q(psu_sn__iexact=clean_sn) |
        Q(monitor_sn__iexact=clean_sn) | Q(avr_sn__iexact=clean_sn) |
        Q(keyboard_sn__iexact=clean_sn) | Q(mouse_sn__iexact=clean_sn)
    )
    if exclude_pc_id:
        pc_qs = pc_qs.exclude(id=exclude_pc_id)
        
    if pc_qs.exists():
        return True
        
    return False

# =========================================================
#  AUTHENTICATION (LOGIN / LOGOUT)
# =========================================================

def custom_login(request):
    if request.user.is_authenticated:
        return redirect('dashboard')

    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            auth_login(request, user)
            if 'next' in request.GET:
                return redirect(request.GET.get('next'))
            return redirect('dashboard')
    else:
        form = AuthenticationForm()
    
    return render(request, 'pages/login.html', {'form': form})

def logout_view(request):
    logout(request)
    try:
        if os.name == 'nt':
            script_path = os.path.join(os.getcwd(), "auto_backup.ps1")
            
            if os.path.exists(script_path):
                subprocess.Popen([
                    "powershell.exe", 
                    "-ExecutionPolicy", "Bypass", 
                    "-WindowStyle", "Hidden", 
                    "-File", script_path
                ])
    except Exception as e:
        print(f"Backup script failed to trigger: {e}")
    return redirect('login')

# =========================================================
#  DASHBOARD & ID GENERATION
# =========================================================

@login_required
def dashboard(request):
    active_count = PersonalComputer.objects.filter(status='Working').count()
    total_count = PersonalComputer.objects.count()
    percentage = int((active_count / total_count) * 100) if total_count > 0 else 0

    rooms_data = []
    for room in Room.objects.all():
        pcs = PersonalComputer.objects.filter(room=room)
        rooms_data.append({'id': room.id, 'name': room.name, 'count': pcs.count(), 'pcs': pcs})

    defective_pcs = PersonalComputer.objects.filter(status='Defective')
    
    incomplete_pcs = PersonalComputer.objects.filter(status='Working').filter(
        # Check core internals
        Q(processor__isnull=True) | Q(processor__exact='') |
        Q(ram__isnull=True) | Q(ram__exact='') |
        Q(storage__isnull=True) | Q(storage__exact='') |
        Q(motherboard__isnull=True) | Q(motherboard__exact='') |
        Q(psu__isnull=True) | Q(psu__exact='') |
        (Q(is_igpu=False) & (Q(graphics_card__isnull=True) | Q(graphics_card__exact=''))) |
        # Check peripherals
        Q(monitor_details__isnull=True) | Q(monitor_details__exact='') |
        Q(avr_details__isnull=True) | Q(avr_details__exact='') |
        Q(keyboard_details__isnull=True) | Q(keyboard_details__exact='') |
        Q(mouse_details__isnull=True) | Q(mouse_details__exact='') |
        (Q(has_hdmi=False) & Q(has_vga=False))
    ).distinct()

    loose_tools = Tool.objects.filter(status='Working', linked_pc__isnull=True)
    
    avail_cpu = loose_tools.filter(category='Computer Part').filter(Q(name__icontains='core') | Q(name__icontains='ryzen') | Q(name__icontains='intel') | Q(name__icontains='amd') | Q(name__icontains='processor') | Q(name__icontains='cpu'))
    avail_ram = loose_tools.filter(category='Computer Part').filter(Q(name__icontains='ram') | Q(name__icontains='gb') | Q(name__icontains='ddr')).exclude(name__icontains='ssd').exclude(name__icontains='hdd').exclude(name__icontains='sata')
    avail_storage = loose_tools.filter(category='Computer Part').filter(Q(name__icontains='hdd') | Q(name__icontains='ssd') | Q(name__icontains='tb') | Q(name__icontains='drive') | Q(name__icontains='sata') | Q(name__icontains='nvme'))
    avail_gpu = loose_tools.filter(category='Computer Part').filter(Q(name__icontains='gtx') | Q(name__icontains='rtx') | Q(name__icontains='rx') | Q(name__icontains='radeon') | Q(name__icontains='gpu') | Q(name__icontains='graphics') | Q(name__icontains='video'))
    avail_mobo = loose_tools.filter(category='Computer Part').filter(Q(name__icontains='motherboard') | Q(name__icontains='mobo') | Q(name__icontains='asrock') | Q(name__icontains='msi') | Q(name__icontains='asus') | Q(name__icontains='gigabyte'))
    avail_psu = loose_tools.filter(category='Computer Part').filter(Q(name__icontains='psu') | Q(name__icontains='power supply') | Q(name__icontains='watts') | Q(name__icontains='corsair') | Q(name__icontains='seasonic') | Q(name__icontains='evga') | Q(name__icontains='thermaltake') | Q(name__icontains='fsp') | Q(name__icontains='silverstone'))
    
    avail_monitors = loose_tools.filter(category__in=['Accessory', 'Network']).filter(Q(name__icontains='monitor') | Q(name__icontains='vision') | Q(name__icontains='viewplus') | Q(name__icontains='samsung') | Q(name__icontains='aoc') | Q(name__icontains='acer') | Q(name__icontains='display'))
    avail_avrs = loose_tools.filter(category__in=['Accessory', 'Network']).filter(Q(name__icontains='avr') | Q(name__icontains='secure') | Q(name__icontains='ups') | Q(name__icontains='voltage'))
    avail_keyboards = loose_tools.filter(category__in=['Accessory', 'Network']).filter(Q(name__icontains='keyboard') | Q(name__icontains='a4tech') | Q(name__icontains='logitech') | Q(name__icontains='inplay') | Q(name__icontains='k0'))
    avail_mice = loose_tools.filter(category__in=['Accessory', 'Network']).filter(Q(name__icontains='mouse') | Q(name__icontains='a4tech') | Q(name__icontains='logitech') | Q(name__icontains='inplay') | Q(name__icontains='m0'))

    context = {
        'active': active_count, 'total': total_count, 'percentage': percentage,
        'rooms': rooms_data, 'defective': defective_pcs, 'incomplete': incomplete_pcs,
        'avail_cpu': avail_cpu, 'avail_ram': avail_ram, 'avail_storage': avail_storage,
        'avail_gpu': avail_gpu, 'avail_mobo': avail_mobo, 'avail_psu': avail_psu, 'avail_monitors': avail_monitors,
        'avail_avrs': avail_avrs, 'avail_keyboards': avail_keyboards, 'avail_mice': avail_mice,
    }
    return render(request, 'pages/dashboard.html', context)

@login_required
def get_next_pc_id(request):
    room_id = request.GET.get('room_id')
    if not room_id:
        return JsonResponse({'next_id': ''})

    room = get_object_or_404(Room, id=room_id)
    prefix = room.name.split()[0].upper()
    existing_ids = PersonalComputer.objects.filter(room=room).values_list('custom_id', flat=True)
    
    max_number = 0
    for pc_id in existing_ids:
        match = re.search(r'(\d+)$', pc_id)
        if match:
            number = int(match.group(1))
            if number > max_number:
                max_number = number
    
    next_number = max_number + 1
    new_id = f"{prefix}-PC-{next_number:02d}"
    
    return JsonResponse({'next_id': new_id})

@login_required
def add_room(request):
    if request.method == "POST":
        room_name = request.POST.get('name')
        if room_name:
            if not Room.objects.filter(name=room_name).exists():
                Room.objects.create(name=room_name)
    return redirect('dashboard')

@login_required
def edit_room(request, room_id):
    room = get_object_or_404(Room, id=room_id)
    if request.method == "POST":
        new_name = request.POST.get('name')
        if new_name:
            room.name = new_name
            room.save()
    return redirect('dashboard')

@login_required
def delete_room(request, room_id):
    room = get_object_or_404(Room, id=room_id)
    room.delete()
    return redirect('dashboard')

# =========================================================
#  PC MANAGEMENT & AUTO-GENERATION
# =========================================================

@login_required
def add_pc(request):
    if request.method == 'POST':
        # 0. COPY THE DATA SO WE CAN SCRUB IT
        data = request.POST.copy()
        
        # --- BACKEND FAILSAFE: No Name = No Serial Number ---
        component_pairs = [
            ('processor', 'processor_sn'), ('ram', 'ram_sn'), ('ram_2', 'ram_2_sn'),
            ('ram_3', 'ram_3_sn'), ('ram_4', 'ram_4_sn'), ('storage', 'storage_sn'),
            ('storage_2', 'storage_2_sn'), ('storage_3', 'storage_3_sn'), ('storage_4', 'storage_4_sn'),
            ('graphics_card', 'graphics_card_sn'), ('motherboard', 'motherboard_sn'),
            ('psu', 'psu_sn'),
            ('monitor_details', 'monitor_sn'), ('avr_details', 'avr_sn'),
            ('keyboard_details', 'keyboard_sn'), ('mouse_details', 'mouse_sn')
        ]
        
        for name_field, sn_field in component_pairs:
            if not data.get(name_field) or data.get(name_field).strip() == '':
                data[sn_field] = ''  # Ruthlessly wipe the S/N if the Name is empty
        # ----------------------------------------------------

        # PRE-CHECK: Validate all Serial Numbers before saving anything
        is_igpu = data.get('is_igpu') == 'true'
        sn_fields = [
            data.get('processor_sn'), data.get('ram_sn'), data.get('ram_2_sn'), data.get('ram_3_sn'), data.get('ram_4_sn'),
            data.get('storage_sn'), data.get('storage_2_sn'), data.get('storage_3_sn'), data.get('storage_4_sn'),
            data.get('motherboard_sn'), data.get('psu_sn'), data.get('monitor_sn'), data.get('avr_sn'), data.get('keyboard_sn'), data.get('mouse_sn')
        ]
        if not is_igpu: sn_fields.append(data.get('graphics_card_sn'))
        
        # Apply the blank slot filter
        clean_sn_list = filter_blank_slots(sn_fields)

        for sn in clean_sn_list:
            # FIX APPLIED HERE: Added allow_loose=True so we can fetch storage items!
            if is_duplicate_sn(sn, allow_loose=True):
                messages.error(request, f"Error: Serial Number '{sn}' is already registered in the system.")
                return redirect('dashboard')
        
        room_id = data.get('room')
        target_room = get_object_or_404(Room, id=room_id)
        has_hdmi = 'has_hdmi' in data
        has_vga = 'has_vga' in data
        
        # 1. CREATE THE PC RECORD WITH MULTIPLE SLOTS & GPU TYPE
        new_pc = PersonalComputer.objects.create(
            custom_id = data.get('custom_id'),
            room = target_room,
            status = data.get('status'),
            
            processor = data.get('processor'), processor_sn = data.get('processor_sn'),
            ram = data.get('ram'), ram_sn = data.get('ram_sn'),
            ram_2 = data.get('ram_2'), ram_2_sn = data.get('ram_2_sn'),
            ram_3 = data.get('ram_3'), ram_3_sn = data.get('ram_3_sn'),
            ram_4 = data.get('ram_4'), ram_4_sn = data.get('ram_4_sn'),
            
            storage = data.get('storage'), storage_sn = data.get('storage_sn'),
            storage_2 = data.get('storage_2'), storage_2_sn = data.get('storage_2_sn'),
            storage_3 = data.get('storage_3'), storage_3_sn = data.get('storage_3_sn'),
            storage_4 = data.get('storage_4'), storage_4_sn = data.get('storage_4_sn'),
            
            is_igpu = is_igpu,
            graphics_card = data.get('graphics_card'),
            graphics_card_sn = data.get('graphics_card_sn') if not is_igpu else "",
            
            motherboard = data.get('motherboard'), motherboard_sn = data.get('motherboard_sn'),
            psu = data.get('psu'), psu_sn = data.get('psu_sn'),
            monitor_details = data.get('monitor_details'), monitor_sn = data.get('monitor_sn'),
            avr_details = data.get('avr_details'), avr_sn = data.get('avr_sn'),
            keyboard_details = data.get('keyboard_details'), keyboard_sn = data.get('keyboard_sn'),
            mouse_details = data.get('mouse_details'), mouse_sn = data.get('mouse_sn'),
            
            has_hdmi = has_hdmi, has_vga = has_vga 
        )

        # 2. SMART ASSET GENERATION
        component_map = [
            (new_pc.processor, new_pc.processor_sn, "Computer Part"),
            (new_pc.ram, new_pc.ram_sn, "Computer Part"), (new_pc.ram_2, new_pc.ram_2_sn, "Computer Part"),
            (new_pc.ram_3, new_pc.ram_3_sn, "Computer Part"), (new_pc.ram_4, new_pc.ram_4_sn, "Computer Part"),
            (new_pc.storage, new_pc.storage_sn, "Computer Part"), (new_pc.storage_2, new_pc.storage_2_sn, "Computer Part"),
            (new_pc.storage_3, new_pc.storage_3_sn, "Computer Part"), (new_pc.storage_4, new_pc.storage_4_sn, "Computer Part"),
            (new_pc.motherboard, new_pc.motherboard_sn, "Computer Part"),
            (new_pc.psu, new_pc.psu_sn, "Computer Part"),
            (new_pc.monitor_details, new_pc.monitor_sn, "Accessory"),
            (new_pc.avr_details, new_pc.avr_sn, "Accessory"),
            (new_pc.keyboard_details, new_pc.keyboard_sn, "Accessory"),
            (new_pc.mouse_details, new_pc.mouse_sn, "Accessory")
        ]
        
        if not is_igpu:
            component_map.append((new_pc.graphics_card, new_pc.graphics_card_sn, "Computer Part"))

        if has_hdmi: component_map.append(("HDMI Cable", "", "Accessory"))
        if has_vga: component_map.append(("VGA Cable", "", "Accessory"))

        for field_name, field_sn, cat in component_map:
            if field_name and field_name.strip():  
                clean_name = field_name.strip()
                clean_sn = field_sn.strip() if field_sn else ""
                
                qs = Tool.objects.filter(name__iexact=clean_name, linked_pc__isnull=True, status='Working', category=cat)
                existing_tool = qs.filter(serial_number__iexact=clean_sn).first() if clean_sn else qs.filter(Q(serial_number__isnull=True) | Q(serial_number__exact='')).first() or qs.first()
                
                if existing_tool:
                    if existing_tool.quantity > 1:
                        existing_tool.quantity -= 1
                        existing_tool.save()
                        Tool.objects.create(name=clean_name, serial_number=clean_sn, category=cat, quantity=1, status="In Use", room=target_room, linked_pc=new_pc)
                    else:
                        existing_tool.room = target_room
                        existing_tool.status = "In Use"
                        existing_tool.linked_pc = new_pc
                        if clean_sn: existing_tool.serial_number = clean_sn
                        existing_tool.save()
                else:
                    Tool.objects.create(name=clean_name, serial_number=clean_sn, category=cat, quantity=1, status="In Use", room=target_room, linked_pc=new_pc)
                    
        return redirect('dashboard')
    return redirect('dashboard')


@login_required
def update_pc(request, pc_id):
    pc = get_object_or_404(PersonalComputer, id=pc_id)
    if request.method == 'POST':
        # 0. COPY THE DATA SO WE CAN SCRUB IT
        data = request.POST.copy()
        
        # --- BACKEND FAILSAFE: No Name = No Serial Number ---
        component_pairs = [
            ('processor', 'processor_sn'), ('ram', 'ram_sn'), ('ram_2', 'ram_2_sn'),
            ('ram_3', 'ram_3_sn'), ('ram_4', 'ram_4_sn'), ('storage', 'storage_sn'),
            ('storage_2', 'storage_2_sn'), ('storage_3', 'storage_3_sn'), ('storage_4', 'storage_4_sn'),
            ('graphics_card', 'graphics_card_sn'), ('motherboard', 'motherboard_sn'),
            ('psu', 'psu_sn'),
            ('monitor_details', 'monitor_sn'), ('avr_details', 'avr_sn'),
            ('keyboard_details', 'keyboard_sn'), ('mouse_details', 'mouse_sn')
        ]
        
        for name_field, sn_field in component_pairs:
            if not data.get(name_field) or data.get(name_field).strip() == '':
                data[sn_field] = ''  # Ruthlessly wipe the S/N if the Name is empty
        # ----------------------------------------------------

        # PRE-CHECK: Validate all Serial Numbers before saving anything
        is_igpu = data.get('is_igpu') == 'true'
        sn_fields = [
            data.get('processor_sn'), data.get('ram_sn'), data.get('ram_2_sn'), data.get('ram_3_sn'), data.get('ram_4_sn'),
            data.get('storage_sn'), data.get('storage_2_sn'), data.get('storage_3_sn'), data.get('storage_4_sn'),
            data.get('motherboard_sn'), data.get('psu_sn'), data.get('monitor_sn'), data.get('avr_sn'), data.get('keyboard_sn'), data.get('mouse_sn')
        ]
        if not is_igpu: sn_fields.append(data.get('graphics_card_sn'))
        
        # Apply the blank slot filter
        clean_sn_list = filter_blank_slots(sn_fields)

        for sn in clean_sn_list:
            # FIX APPLIED HERE: Added allow_loose=True so we can fetch storage items!
            if is_duplicate_sn(sn, exclude_pc_id=pc.id, allow_loose=True):
                messages.error(request, f"Error: Serial Number '{sn}' is already registered in the system.")
                return redirect('dashboard')
        
        room_id = data.get('room')
        if room_id:
            target_room = get_object_or_404(Room, id=room_id)
            pc.room = target_room
        else:
            target_room = None; pc.room = None
            
        pc.custom_id = data.get('custom_id')
        pc.status = data.get('status')
        pc.remarks = data.get('remarks')
        
        # 1. CAPTURE OLD SPECS
        old_specs = {
            'processor': (pc.processor, pc.processor_sn),
            'ram': (pc.ram, pc.ram_sn), 'ram_2': (pc.ram_2, pc.ram_2_sn), 'ram_3': (pc.ram_3, pc.ram_3_sn), 'ram_4': (pc.ram_4, pc.ram_4_sn),
            'storage': (pc.storage, pc.storage_sn), 'storage_2': (pc.storage_2, pc.storage_2_sn), 'storage_3': (pc.storage_3, pc.storage_3_sn), 'storage_4': (pc.storage_4, pc.storage_4_sn),
            'motherboard': (pc.motherboard, pc.motherboard_sn),
            'psu': (pc.psu, pc.psu_sn),
            'monitor_details': (pc.monitor_details, pc.monitor_sn), 'avr_details': (pc.avr_details, pc.avr_sn),
            'keyboard_details': (pc.keyboard_details, pc.keyboard_sn), 'mouse_details': (pc.mouse_details, pc.mouse_sn),
        }
        
        old_igpu = pc.is_igpu
        old_gpu = pc.graphics_card
        if not old_igpu and old_gpu:
            old_specs['graphics_card'] = (old_gpu, pc.graphics_card_sn)
            
        old_has_hdmi = pc.has_hdmi; old_has_vga = pc.has_vga
        
        # SAVE NEW SPECS
        pc.is_igpu = is_igpu
        pc.processor = data.get('processor'); pc.processor_sn = data.get('processor_sn')
        pc.ram = data.get('ram'); pc.ram_sn = data.get('ram_sn')
        pc.ram_2 = data.get('ram_2'); pc.ram_2_sn = data.get('ram_2_sn')
        pc.ram_3 = data.get('ram_3'); pc.ram_3_sn = data.get('ram_3_sn')
        pc.ram_4 = data.get('ram_4'); pc.ram_4_sn = data.get('ram_4_sn')
        pc.storage = data.get('storage'); pc.storage_sn = data.get('storage_sn')
        pc.storage_2 = data.get('storage_2'); pc.storage_2_sn = data.get('storage_2_sn')
        pc.storage_3 = data.get('storage_3'); pc.storage_3_sn = data.get('storage_3_sn')
        pc.storage_4 = data.get('storage_4'); pc.storage_4_sn = data.get('storage_4_sn')
        pc.graphics_card = data.get('graphics_card'); pc.graphics_card_sn = data.get('graphics_card_sn') if not pc.is_igpu else ""
        pc.motherboard = data.get('motherboard'); pc.motherboard_sn = data.get('motherboard_sn')
        pc.psu = data.get('psu'); pc.psu_sn = data.get('psu_sn')
        pc.monitor_details = data.get('monitor_details'); pc.monitor_sn = data.get('monitor_sn') 
        pc.avr_details = data.get('avr_details'); pc.avr_sn = data.get('avr_sn') 
        pc.keyboard_details = data.get('keyboard_details'); pc.keyboard_sn = data.get('keyboard_sn') 
        pc.mouse_details = data.get('mouse_details'); pc.mouse_sn = data.get('mouse_sn')       
        pc.has_hdmi = 'has_hdmi' in data; pc.has_vga = 'has_vga' in data 
        
        # 2. CLEANUP OLD REMOVED PARTS
        for field, (old_name, old_sn) in old_specs.items():
            new_name = data.get(field, '').strip() if not (field == 'graphics_card' and pc.is_igpu) else ''
            new_sn = data.get(f"{field}_sn", '').strip()
            old_name = old_name.strip() if old_name else ""
            old_sn = old_sn.strip() if old_sn else ""
            
            if old_name and (old_name != new_name or old_sn != new_sn):
                action_choice = data.get(f'action_{field}', 'Working')
                old_tool = Tool.objects.filter(linked_pc=pc, name__iexact=old_name).first()
                if old_tool:
                    if action_choice in ['Defective', 'Missing']:
                        old_tool.linked_pc = None; old_tool.room = None; old_tool.status = action_choice; old_tool.save()
                    else: 
                        qs = Tool.objects.filter(name__iexact=old_tool.name, category=old_tool.category, room__isnull=True, linked_pc__isnull=True, status='Working')
                        existing_loose = qs.filter(serial_number__iexact=old_tool.serial_number).first() if old_tool.serial_number else qs.filter(Q(serial_number__isnull=True) | Q(serial_number__exact='')).first()
                        if existing_loose:
                            existing_loose.quantity += old_tool.quantity; existing_loose.save(); old_tool.delete() 
                        else:
                            old_tool.linked_pc = None; old_tool.room = None; old_tool.status = 'Working'; old_tool.save()

        for cab_name, old_check, new_check in [("HDMI Cable", old_has_hdmi, pc.has_hdmi), ("VGA Cable", old_has_vga, pc.has_vga)]:
            if old_check and not new_check:
                old_cab = Tool.objects.filter(linked_pc=pc, name__iexact=cab_name).first()
                if old_cab:
                    existing = Tool.objects.filter(name__iexact=cab_name, room__isnull=True, linked_pc__isnull=True, status='Working').first()
                    if existing: existing.quantity += old_cab.quantity; existing.save(); old_cab.delete()
                    else: old_cab.linked_pc = None; old_cab.room = None; old_cab.status = 'Working'; old_cab.save()

        # 3. FETCH/CREATE NEW COMPONENTS
        new_components = [
            (pc.processor, pc.processor_sn, "Computer Part"), (pc.ram, pc.ram_sn, "Computer Part"),
            (pc.ram_2, pc.ram_2_sn, "Computer Part"), (pc.ram_3, pc.ram_3_sn, "Computer Part"), (pc.ram_4, pc.ram_4_sn, "Computer Part"),
            (pc.storage, pc.storage_sn, "Computer Part"), (pc.storage_2, pc.storage_2_sn, "Computer Part"),
            (pc.storage_3, pc.storage_3_sn, "Computer Part"), (pc.storage_4, pc.storage_4_sn, "Computer Part"),
            (pc.motherboard, pc.motherboard_sn, "Computer Part"), (pc.psu, pc.psu_sn, "Computer Part"), 
            (pc.monitor_details, pc.monitor_sn, "Accessory"), (pc.avr_details, pc.avr_sn, "Accessory"), 
            (pc.keyboard_details, pc.keyboard_sn, "Accessory"), (pc.mouse_details, pc.mouse_sn, "Accessory")
        ]
        if not pc.is_igpu: new_components.append((pc.graphics_card, pc.graphics_card_sn, "Computer Part"))
        if pc.has_hdmi: new_components.append(("HDMI Cable", "", "Accessory"))
        if pc.has_vga: new_components.append(("VGA Cable", "", "Accessory"))
        
        attached_tools = Tool.objects.filter(linked_pc=pc)
        for field_name, field_sn, cat in new_components:
            if field_name and field_name.strip():
                clean_name = field_name.strip()
                clean_sn = field_sn.strip() if field_sn else ""
                
                already_attached = False
                for t in attached_tools:
                    t_sn = t.serial_number.strip() if t.serial_number else ""
                    if t.name.lower() == clean_name.lower() and t_sn == clean_sn:
                        already_attached = True
                        break
                        
                if already_attached: continue
                    
                qs = Tool.objects.filter(name__iexact=clean_name, linked_pc__isnull=True, status='Working', category=cat)
                existing_tool = qs.filter(serial_number__iexact=clean_sn).first() if clean_sn else qs.filter(Q(serial_number__isnull=True) | Q(serial_number__exact='')).first() or qs.first()
                
                if existing_tool:
                    if existing_tool.quantity > 1:
                        existing_tool.quantity -= 1; existing_tool.save()
                        Tool.objects.create(name=clean_name, serial_number=clean_sn, category=cat, quantity=1, status="In Use", room=target_room, linked_pc=pc)
                    else:
                        existing_tool.room = target_room; existing_tool.status = "In Use"; existing_tool.linked_pc = pc
                        if clean_sn: existing_tool.serial_number = clean_sn
                        existing_tool.save()
                else:
                    Tool.objects.create(name=clean_name, serial_number=clean_sn, category=cat, quantity=1, status="In Use", room=target_room, linked_pc=pc)

        pc.save()
        return redirect('dashboard')
    
@login_required
def delete_pc(request, pc_id):
    pc = get_object_or_404(PersonalComputer, id=pc_id)
    
    linked_components = Tool.objects.filter(linked_pc=pc)
    for comp in linked_components:
        qs = Tool.objects.filter(
            name__iexact=comp.name, category=comp.category, 
            room__isnull=True, linked_pc__isnull=True, status='Working'
        )
        if comp.serial_number:
            existing_loose_tool = qs.filter(serial_number__iexact=comp.serial_number).first()
        else:
            existing_loose_tool = qs.filter(Q(serial_number__isnull=True) | Q(serial_number__exact='')).first()

        if existing_loose_tool:
            existing_loose_tool.quantity += comp.quantity
            existing_loose_tool.save()
            comp.delete()
        else:
            comp.linked_pc = None  
            comp.status = 'Working' 
            comp.room = None 
            comp.save()
            
    pc.delete()
    return redirect('dashboard')


@login_required
def edit_pc_room(request, pc_id):
    if request.method == 'POST':
        pc = get_object_or_404(PersonalComputer, id=pc_id)
        room_id = request.POST.get('room')
        
        if room_id:
            new_room = get_object_or_404(Room, id=room_id)
            pc.room = new_room
        else:
            pc.room = None
            
        pc.save()

        linked_components = Tool.objects.filter(linked_pc=pc)
        for comp in linked_components:
            comp.room = pc.room
            comp.save()

    return redirect('master_list')

@login_required
def room_list(request, room_id):
    room = get_object_or_404(Room, id=room_id)
    pcs = PersonalComputer.objects.filter(room=room)
    return render(request, 'pages/room_list.html', {'room': room, 'pcs': pcs})

# =========================================================
#  MASTER LIST & BORROWING SYSTEM
# =========================================================

@login_required
def master_list(request):
    all_pcs = PersonalComputer.objects.all()
    all_laptops = Laptop.objects.all()
    all_tools = Tool.objects.all()
    all_rooms = Room.objects.all() 

    avail_pcs = all_pcs.filter(status='Working', is_borrowed=False)
    avail_laptops = all_laptops.filter(status='Working', is_borrowed=False)
    
    borrowable_categories = ['Hand Tool', 'Accessory', 'Network', 'Animation', 'Other']
    
    avail_tools = all_tools.filter(
        status='Working', 
        quantity__gt=0, 
        category__in=borrowable_categories,
        linked_pc__isnull=True 
    )

    borrowed_pcs = all_pcs.filter(is_borrowed=True)
    borrowed_laptops = all_laptops.filter(is_borrowed=True)
    borrowed_tools_list = BorrowedItem.objects.all().select_related('tool_asset')

    # --- ADDED THIS LINE: Filter for Consumables & Office Supplies ---
    consumables = all_tools.filter(category__in=['Office Supplies', 'Consumable', 'Medicines'])

    transaction_list = TransactionHistory.objects.all().order_by('-timestamp')
    paginator = Paginator(transaction_list, 50) 
    page_number = request.GET.get('vault_page')
    transactions = paginator.get_page(page_number)
    print_batch_id = request.session.pop('print_batch_id', None)

    context = {
        'all_rooms': all_rooms,
        'all_pcs': all_pcs, 
        'all_laptops': all_laptops, 
        'all_tools': all_tools,
        'avail_pcs': avail_pcs, 
        'avail_laptops': avail_laptops, 
        'avail_tools': avail_tools, 
        'borrowed_pcs': borrowed_pcs, 
        'borrowed_laptops': borrowed_laptops, 
        'borrowed_tools': borrowed_tools_list, 
        'consumables': consumables, # <--- ADDED TO CONTEXT
        'transactions': transactions, 
        'print_batch_id': print_batch_id, 
    }
    return render(request, 'pages/master_list.html', context)

@login_required
def borrow_item(request):
    return redirect('master_list')

@login_required
def batch_borrow(request):
    if request.method == 'POST':
        selected_items_str = request.POST.get('selected_items')
        b_location = request.POST.get('borrow_location') 
        r_date = request.POST.get('return_date') or None
        
        if b_location == 'In-Campus':
            b_name = request.POST.get('borrower_name_room') 
            b_contact = "N/A"
        else:
            b_name = request.POST.get('borrower_name') 
            b_contact = request.POST.get('borrower_contact')
        
        b_position = request.POST.get('borrower_position')
        b_branch = request.POST.get('borrower_branch')
        b_email = request.POST.get('borrower_email')
        b_purpose = request.POST.get('borrower_purpose')
            
        new_batch_id = str(uuid.uuid4())

        if selected_items_str:
            item_list = selected_items_str.split(',')
            for item_str in item_list:
                if ':' in item_str:
                    itype, iid = item_str.split(':')
                    
                    if itype == 'pc': 
                        obj = get_object_or_404(PersonalComputer, id=iid)
                        obj.is_borrowed = True
                        obj.borrower_name = b_name
                        obj.borrower_contact = b_contact
                        obj.borrow_location = b_location
                        obj.return_date = r_date
                        obj.batch_id = new_batch_id
                        obj.borrower_position = b_position
                        obj.borrower_branch = b_branch
                        obj.borrower_email = b_email
                        obj.borrower_purpose = b_purpose
                        obj.save()

                        TransactionHistory.objects.create(
                            action='Borrow',
                            item_details=f"PC Unit ({obj.custom_id})",
                            borrower_name=b_name,
                            batch_id=new_batch_id
                        )
                    
                    elif itype == 'laptop': 
                        obj = get_object_or_404(Laptop, id=iid)
                        obj.is_borrowed = True
                        obj.borrower_name = b_name
                        obj.borrower_contact = b_contact
                        obj.borrow_location = b_location
                        obj.return_date = r_date
                        obj.batch_id = new_batch_id
                        obj.borrower_position = b_position
                        obj.borrower_branch = b_branch
                        obj.borrower_email = b_email
                        obj.borrower_purpose = b_purpose
                        obj.save()

                        TransactionHistory.objects.create(
                            action='Borrow',
                            item_details=f"Laptop {obj.brand} ({obj.custom_id})",
                            borrower_name=b_name,
                            batch_id=new_batch_id
                        )
                    
                    elif itype == 'tool': 
                        obj = get_object_or_404(Tool, id=iid)
                        qty_key = f"quantity_tool_{iid}"
                        qty_to_borrow = int(request.POST.get(qty_key, 1))

                        if obj.quantity >= qty_to_borrow and qty_to_borrow > 0:
                            obj.quantity -= qty_to_borrow
                            obj.save()
                            
                            BorrowedItem.objects.create(
                                tool_asset=obj,
                                borrower_name=b_name,
                                borrower_contact=b_contact,
                                borrow_location=b_location,
                                quantity_borrowed=qty_to_borrow, 
                                return_due_date=r_date,
                                batch_id=new_batch_id,
                                borrower_position=b_position,
                                borrower_branch=b_branch,
                                borrower_email=b_email,
                                borrower_purpose=b_purpose
                            )

                            TransactionHistory.objects.create(
                                action='Borrow',
                                item_details=f"{obj.name} (Qty: {qty_to_borrow})",
                                borrower_name=b_name,
                                batch_id=new_batch_id
                            )
            
            if b_location in ['Transfer', 'Outside']:
                request.session['print_batch_id'] = new_batch_id

    return redirect('master_list')

@login_required
def return_item(request, item_type, item_id):
    if item_type == 'pc': 
        obj = get_object_or_404(PersonalComputer, id=item_id)
        old_borrower = obj.borrower_name
        obj.is_borrowed = False; obj.borrower_name = None; obj.borrower_contact = None; obj.borrow_location = None; obj.return_date = None; obj.batch_id = None
        obj.save()
        TransactionHistory.objects.create(action='Return', item_details=f"PC Unit ({obj.custom_id})", borrower_name=old_borrower or "Unknown")
        
    elif item_type == 'laptop': 
        obj = get_object_or_404(Laptop, id=item_id)
        old_borrower = obj.borrower_name
        obj.is_borrowed = False; obj.borrower_name = None; obj.borrower_contact = None; obj.borrow_location = None; obj.return_date = None; obj.batch_id = None
        obj.save()
        TransactionHistory.objects.create(action='Return', item_details=f"Laptop {obj.brand} ({obj.custom_id})", borrower_name=old_borrower or "Unknown")
        
    elif item_type == 'borrowed_tool': 
        transaction = get_object_or_404(BorrowedItem, id=item_id)
        tool = transaction.tool_asset
        old_borrower = transaction.borrower_name
        tool.quantity += transaction.quantity_borrowed
        tool.save()
        TransactionHistory.objects.create(action='Return', item_details=f"{tool.name} (Qty: {transaction.quantity_borrowed})", borrower_name=old_borrower or "Unknown")
        transaction.delete()
        
    return redirect('master_list')

@login_required
def batch_return(request):
    if request.method == 'POST':
        selected_items_str = request.POST.get('selected_items')
        if selected_items_str:
            item_list = selected_items_str.split(',')
            for item_str in item_list:
                if ':' in item_str:
                    itype, iid = item_str.split(':')
                    
                    qty_key = f"return_qty_{itype}_{iid}"
                    return_qty = int(request.POST.get(qty_key, 1))

                    if itype == 'pc': 
                        obj = get_object_or_404(PersonalComputer, id=iid)
                        old_borrower = obj.borrower_name 
                        old_batch = obj.batch_id
                        obj.is_borrowed = False; obj.borrower_name = None; obj.borrower_contact = None; obj.borrow_location = None; obj.return_date = None; obj.batch_id = None
                        obj.save()
                        
                        TransactionHistory.objects.create(
                            action='Return',
                            item_details=f"PC Unit ({obj.custom_id})",
                            borrower_name=old_borrower or "Unknown",
                            batch_id=old_batch
                        )
                        
                    elif itype == 'laptop': 
                        obj = get_object_or_404(Laptop, id=iid)
                        old_borrower = obj.borrower_name 
                        old_batch = obj.batch_id
                        obj.is_borrowed = False; obj.borrower_name = None; obj.borrower_contact = None; obj.borrow_location = None; obj.return_date = None; obj.batch_id = None
                        obj.save()

                        TransactionHistory.objects.create(
                            action='Return',
                            item_details=f"Laptop {obj.brand} ({obj.custom_id})",
                            borrower_name=old_borrower or "Unknown",
                            batch_id=old_batch
                        )

                    elif itype == 'borrowed_tool':
                        transaction = get_object_or_404(BorrowedItem, id=iid)
                        tool = transaction.tool_asset
                        
                        old_borrower = transaction.borrower_name
                        old_batch = transaction.batch_id

                        actual_return_qty = min(return_qty, transaction.quantity_borrowed)

                        tool.quantity += actual_return_qty
                        tool.save()
                        
                        TransactionHistory.objects.create(
                            action='Return',
                            item_details=f"{tool.name} (Qty: {actual_return_qty})",
                            borrower_name=old_borrower or "Unknown",
                            batch_id=old_batch
                        )
                        
                        if actual_return_qty >= transaction.quantity_borrowed:
                            transaction.delete()
                        else:
                            transaction.quantity_borrowed -= actual_return_qty
                            transaction.save()

    return redirect('master_list')

# =========================================================
#  DOCUMENT GENERATION
# =========================================================

def generate_word_doc(request, batch_id):
    pcs = PersonalComputer.objects.filter(batch_id=batch_id)
    laptops = Laptop.objects.filter(batch_id=batch_id)
    borrowed_tools = BorrowedItem.objects.filter(batch_id=batch_id)
    
    item_list = []
    def add_to_list(qty, name, status):
        item_list.append({'qty': qty, 'name': name, 'status': status})

    for pc in pcs: add_to_list("1", f"PC Unit ({pc.custom_id})", pc.status)
    for lap in laptops: add_to_list("1", f"Laptop {lap.brand} ({lap.custom_id})", lap.status)
    for t in borrowed_tools: add_to_list(str(t.quantity_borrowed), f"{t.tool_asset.name}", t.tool_asset.status)

    if not item_list:
        return redirect('master_list')

    first = pcs.first() or laptops.first() or borrowed_tools.first()
    if not first: 
        return redirect('master_list')

    borrower_name = first.borrower_name
    location = first.borrow_location

    template_path = os.path.join(settings.BASE_DIR, 'pages', 'templates', 'docx', 'transmittal_template.docx')
    doc = DocxTemplate(template_path)

    context = {
        'today_date': timezone.now().strftime("%B %d, %Y"),
        'destination': f"{location} - {borrower_name}",
        'delivered_by': "Inventory Custodian",
    }
    doc.render(context)

    target_table = doc.tables[1]
    for item in item_list:
        new_row = target_table.add_row()
        row_values = [str(item['qty']), str(item['name']), str(item['status'])]
        
        for index, value in enumerate(row_values):
            cell = new_row.cells[index]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)

    raw_name = (first.borrower_name or "Unknown").strip().replace(" ", "_")
    safe_name = re.sub(r'[\\/*?:"<>|]', "", raw_name)
    output_filename = f'Transmittal_{safe_name}.docx'    

    downloads_folder = str(Path.home() / "Downloads")
    output_filepath = os.path.join(downloads_folder, output_filename)

    try:
        doc.save(output_filepath)
    except PermissionError:
        import time
        output_filename = f'Transmittal_{safe_name}_{int(time.time())}.docx'
        output_filepath = os.path.join(downloads_folder, output_filename)
        doc.save(output_filepath)

    try:
        if os.name == 'nt': 
            subprocess.Popen(f'explorer /select,"{output_filepath}"')
        else:
            os.startfile(downloads_folder)
    except Exception as e:
        print(f"Failed to open file location: {e}")

    return redirect('master_list')

def reprint_receipt(request, batch_id):
    asset_list = []
    pcs = PersonalComputer.objects.filter(batch_id=batch_id)
    laptops = Laptop.objects.filter(batch_id=batch_id)
    tools = BorrowedItem.objects.filter(batch_id=batch_id) 

    for obj in pcs: asset_list.append({'name': f"PC ({obj.custom_id})", 'category': "PC", 'status': obj.status, 'custom_id': obj.custom_id, 'obj': obj})
    for obj in laptops: asset_list.append({'name': f"{obj.brand} {obj.model}", 'category': "Laptop", 'status': obj.status, 'custom_id': obj.custom_id, 'obj': obj})
    for t in tools: asset_list.append({'name': t.tool_asset.name, 'category': t.tool_asset.category, 'status': t.tool_asset.status, 'custom_id': f"Qty: {t.quantity_borrowed}", 'obj': t})

    if not asset_list: return redirect('master_list')

    borrower_info = asset_list[0]['obj']

    if borrower_info.borrow_location == 'Transfer':
        return redirect('generate_word_doc', batch_id=batch_id)
    elif borrower_info.borrow_location == 'Outside':
        return redirect('generate_borrower_slip', batch_id=batch_id)

    context = { 'asset_list': asset_list, 'borrower': borrower_info, 'today': timezone.now() }
    return render(request, 'pages/print_receipt.html', context)

# =========================================================
#  TOOL & OFFICE SUPPLIES MANAGEMENT CRUD
# =========================================================

@login_required
def manage_tools(request):
    rooms = Room.objects.all()
    
    # SMART GROUPING FUNCTION: Stacks identical names, sums QTY, retains S/Ns
    def group_tools(qs):
        grouped = {}
        for t in qs:
            if t.name not in grouped:
                grouped[t.name] = {'name': t.name, 'category': t.category, 'total_qty': 0, 'items': []}
            grouped[t.name]['total_qty'] += t.quantity
            grouped[t.name]['items'].append(t)
        # Returns a sorted list of grouped dictionaries
        return sorted(grouped.values(), key=lambda x: x['name'].lower())

    # 1. Fetch Queries
    unassigned_qs = Tool.objects.filter(room__isnull=True, status='Working').exclude(category__in=['Office Supplies', 'Medicines', 'Consumable'])
    defective_qs = Tool.objects.filter(status='Defective').exclude(category__in=['Office Supplies', 'Medicines', 'Consumable'])
    
    # 2. Group General & Defective
    unassigned_tools = group_tools(unassigned_qs)
    defective_tools = group_tools(defective_qs)
    
    # 3. Pre-group tools for every specific Room
    for room in rooms:
        room_tools_qs = room.tool_set.filter(linked_pc__isnull=True).exclude(category__in=['Office Supplies', 'Medicines', 'Consumable', 'Computer Part'])
        room.grouped_tools = group_tools(room_tools_qs)
        
        room_loose_qs = room.tool_set.filter(linked_pc__isnull=True, category='Computer Part')
        room.grouped_loose = group_tools(room_loose_qs)
        
    context = {'rooms': rooms, 'unassigned_tools': unassigned_tools, 'defective_tools': defective_tools}
    return render(request, 'pages/manage_tools.html', context)

@login_required
def manage_office_supplies(request):
    rooms = Room.objects.all()
    supplies = Tool.objects.filter(category__in=['Office Supplies', 'Medicines', 'Consumable'])
    context = {'rooms': rooms, 'unassigned_tools': supplies} 
    return render(request, 'pages/manage_office_supplies.html', context)

@login_required
def add_tool(request):
    if request.method == 'POST':
        data = request.POST.copy()
        if 'status' not in data: data['status'] = 'Working'

        room_id = request.POST.get('room')
        target_room = get_object_or_404(Room, id=room_id) if room_id else None
        
        # 1. Capture the data directly
        name = request.POST.get('name', '').strip()
        category = request.POST.get('category', '')
        status = data.get('status')
        qty = int(request.POST.get('quantity', 1))
        
        # NEW UOM FIELD
        uom = request.POST.get('unit_of_measure', 'pcs')
        
        # 2. Fetch the dynamic array of serial numbers
        sn_list = request.POST.getlist('serial_numbers[]')
        sn_list = [sn.strip() for sn in sn_list if sn.strip()] # Remove empties
        
        # PRE-CHECK: Validate S/N Uniqueness
        for sn in sn_list:
            if is_duplicate_sn(sn):
                messages.error(request, f"Error: Serial Number '{sn}' is already registered.")
                referer = request.META.get('HTTP_REFERER', '')
                return redirect('manage_office_supplies' if 'manage-office-supplies' in referer else 'manage_tools')

        if sn_list:
            # 3a. Create individual items for every Serial Number provided
            for sn in sn_list:
                Tool.objects.create(name=name, category=category, quantity=1, status=status, room=target_room, serial_number=sn, unit_of_measure=uom)
            
            # 3b. If they typed Qty: 5, but only gave 3 Serial Numbers, lump the remaining 2 into a bulk item
            remainder = qty - len(sn_list)
            if remainder > 0:
                Tool.objects.create(name=name, category=category, quantity=remainder, status=status, room=target_room, unit_of_measure=uom)
        else:
            # 3c. No serial numbers provided at all, just do a normal bulk save
            Tool.objects.create(name=name, category=category, quantity=qty, status=status, room=target_room, unit_of_measure=uom)
            
    referer = request.META.get('HTTP_REFERER', '')
    if 'manage-office-supplies' in referer: return redirect('manage_office_supplies')
    return redirect('manage_tools')

@login_required
def edit_tool(request, tool_id):
    tool = get_object_or_404(Tool, id=tool_id)
    if request.method == 'POST':
        # PRE-CHECK: Validate S/N Uniqueness
        sn_val = request.POST.get('serial_number')
        if sn_val and sn_val.strip() and is_duplicate_sn(sn_val, exclude_tool_id=tool.id):
            messages.error(request, f"Error: Serial Number '{sn_val.strip()}' is already registered.")
            referer = request.META.get('HTTP_REFERER', '')
            return redirect('manage_office_supplies' if 'manage-office-supplies' in referer else 'manage_tools')

        # 1. Grab data directly from the POST request (bypassing ToolForm)
        tool.name = request.POST.get('name', tool.name).strip()
        tool.category = request.POST.get('category', tool.category)
        tool.status = request.POST.get('status', 'Working')
        
        # NEW UOM FIELD
        tool.unit_of_measure = request.POST.get('unit_of_measure', tool.unit_of_measure)
        
        try:
            tool.quantity = int(request.POST.get('quantity', tool.quantity))
        except ValueError:
            pass # Keep old quantity if they typed something weird
            
        # 2. Handle S/N Injection
        tool.serial_number = sn_val.strip() if sn_val and sn_val.strip() else None
        
        # 3. WORKFLOW RULE ENFORCEMENT: 
        # If an item has a Serial Number, it MUST be a singular unique physical item.
        if tool.serial_number and tool.quantity > 1:
            tool.quantity = 1
            
        # 4. Handle Room Assignment
        room_id = request.POST.get('room')
        tool.room = get_object_or_404(Room, id=room_id) if room_id else None
        
        tool.save()
            
    referer = request.META.get('HTTP_REFERER', '')
    if 'manage-office-supplies' in referer: return redirect('manage_office_supplies')
    return redirect('manage_tools')

@login_required
def delete_tool(request, tool_id):
    tool = get_object_or_404(Tool, id=tool_id)
    tool.delete()
    referer = request.META.get('HTTP_REFERER', '')
    if 'manage-office-supplies' in referer: return redirect('manage_office_supplies')
    return redirect('manage_tools')

@login_required
def manage_laptops(request):
    laptops = Laptop.objects.all().order_by('custom_id')
    total_laptops = laptops.count()
    working_laptops = laptops.filter(status='Working').count()
    borrowed_laptops = laptops.filter(is_borrowed=True).count()
    context = {'laptops': laptops, 'total': total_laptops, 'working': working_laptops, 'borrowed': borrowed_laptops}
    return render(request, 'pages/manage_laptops.html', context)

@login_required
def add_laptop(request):
    if request.method == 'POST':
        data = request.POST.copy()
        if not data.get('model') or data.get('model').strip() == '': data['model'] = 'N/A'
        if not data.get('serial_number') or data.get('serial_number').strip() == '': data['serial_number'] = 'N/A'
            
        form = LaptopForm(data)
        if form.is_valid(): form.save()
            
    return redirect('manage_laptops')

@login_required
def edit_laptop(request, laptop_id):
    laptop = get_object_or_404(Laptop, id=laptop_id)
    if request.method == 'POST':
        data = request.POST.copy()
        if not data.get('model') or data.get('model').strip() == '': data['model'] = 'N/A'
        if not data.get('serial_number') or data.get('serial_number').strip() == '': data['serial_number'] = 'N/A'
            
        form = LaptopForm(data, instance=laptop)
        if form.is_valid(): form.save()
            
    return redirect('manage_laptops')

@login_required
def delete_laptop(request, laptop_id):
    laptop = get_object_or_404(Laptop, id=laptop_id)
    laptop.delete()
    return redirect('manage_laptops')

@login_required
def get_next_laptop_id(request):
    existing_ids = Laptop.objects.filter(custom_id__startswith='HCCDD_LAPTOP_').values_list('custom_id', flat=True)
    max_number = 0
    for lap_id in existing_ids:
        match = re.search(r'(\d+)$', lap_id)
        if match:
            number = int(match.group(1))
            if number > max_number: max_number = number
    next_number = max_number + 1
    new_id = f"HCCDD_LAPTOP_{next_number:03d}"
    return JsonResponse({'next_id': new_id})

def generate_borrower_slip(request, batch_id):
    pcs = PersonalComputer.objects.filter(batch_id=batch_id)
    laptops = Laptop.objects.filter(batch_id=batch_id)
    borrowed_tools = BorrowedItem.objects.filter(batch_id=batch_id).exclude(tool_asset__name__icontains="Name Plate")
    
    first = pcs.first() or laptops.first() or borrowed_tools.first()
    if not first: return redirect('master_list')

    def fmt_date(d): return d.strftime("%m/%d/%Y") if d else "TBD"

    doc_purpose = first.borrower_purpose or "Official Use"

    asset_list = []
    for pc in pcs: asset_list.append({'qty': '1', 'desc': 'PC Unit', 'serial': pc.custom_id, 'purpose': doc_purpose, 'date': fmt_date(pc.return_date)})
    for lap in laptops: asset_list.append({'qty': '1', 'desc': f"Laptop {lap.brand} {lap.model}", 'serial': lap.custom_id, 'purpose': doc_purpose, 'date': fmt_date(lap.return_date)})
    for t in borrowed_tools: asset_list.append({'qty': str(t.quantity_borrowed), 'desc': t.tool_asset.name, 'serial': 'N/A', 'purpose': doc_purpose, 'date': fmt_date(t.return_due_date)})

    template_path = os.path.join(settings.BASE_DIR, 'pages', 'templates', 'docx', 'borrower_slip_template.docx')
    doc = DocxTemplate(template_path)
    
    context = {
        'borrower_name': (first.borrower_name or "").upper(),
        'date': timezone.now().strftime("%B %d, %Y"),
        'position': first.borrower_position or "",
        'contact': first.borrower_contact or "",
        'branch': first.borrower_branch or first.borrow_location,
        'email': first.borrower_email or "",
    }
    doc.render(context)

    for table in doc.tables:
        if table.rows and len(table.columns) >= 2 and "Description" in table.cell(0, 1).text:
            for row_num, asset in enumerate(asset_list, start=1):
                new_row = table.add_row()
                number_col = str(asset['serial']) if asset['serial'] != 'N/A' else f"Qty: {asset['qty']}"
                row_values = [str(row_num), str(asset['desc']), number_col, "", str(asset['purpose']), str(asset['date'])]
                for index, value in enumerate(row_values):
                    cell = new_row.cells[index]
                    cell.text = value
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(6)

    raw_name = (first.borrower_name or "Unknown").strip().replace(" ", "_")
    safe_name = re.sub(r'[\\/*?:"<>|]', "", raw_name)
    output_filename = f'BorrowerSlip_{safe_name}.docx'    
    
    downloads_folder = str(Path.home() / "Downloads")
    output_filepath = os.path.join(downloads_folder, output_filename)

    try:
        doc.save(output_filepath)
    except PermissionError:
        import time
        output_filename = f'BorrowerSlip_{safe_name}_{int(time.time())}.docx'
        output_filepath = os.path.join(downloads_folder, output_filename)
        doc.save(output_filepath)

    try:
        if os.name == 'nt': 
            subprocess.Popen(f'explorer /select,"{output_filepath}"')
        else:
            os.startfile(downloads_folder)
    except Exception as e:
        print(f"Failed to open file location: {e}")
    
    return redirect('master_list')

@login_required
def reactivate_tool(request, tool_id):
    tool = get_object_or_404(Tool, id=tool_id)
    
    # Resurrect the item: Set it to Working and return to general storage
    tool.status = 'Working'
    tool.room = None
    tool.linked_pc = None
    tool.save()
    
    messages.success(request, f"Success! {tool.name} has been marked as found and returned to storage.")
    
    # Redirect back to where they clicked it
    referer = request.META.get('HTTP_REFERER', '')
    if 'manage-tools' in referer:
        return redirect('manage_tools')
    return redirect('master_list')